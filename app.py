# from flask import Flask, render_template, request, jsonify, redirect, url_for, flash
# from pypdf import PdfReader
# from docx import Document
# import openai
# import json
# import win32com.client  # Only needed if running on Windows for .doc files
# import pythoncom  # For COM initialization
# import tempfile
# import os
# from flask_sqlalchemy import SQLAlchemy
# from werkzeug.security import generate_password_hash, check_password_hash

# app = Flask(__name__)

# # OpenAI API key
# openai.api_key = 'sk-proj-AKzJAY8kXnNv4u4Vb5jhT3BlbkFJNJqFbDIEh06QWkbsRwGs'

# # Flask configuration
# app.config['SECRET_KEY'] = 'your_secret_key'
# app.config['SQLALCHEMY_DATABASE_URI'] = 'postgresql://postgres:pit@localhost/parserlogin'
# app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# # Initialize SQLAlchemy
# db = SQLAlchemy(app)

# class User(db.Model):
#     id = db.Column(db.Integer, primary_key=True)
#     full_name = db.Column(db.String(150), nullable=False)
#     email = db.Column(db.String(150), unique=True, nullable=False)
#     password = db.Column(db.String(150), nullable=False)

#     def __init__(self, full_name, email, password):
#         self.full_name = full_name
#         self.email = email
#         self.password = generate_password_hash(password, method='pbkdf2:sha256')

# # ats_extractor function using the OpenAI API
# def ats_extractor(resume_data):
#     prompt = '''
#     You are an expert in resume parsing. Please extract all relevant information from the provided resume file. Make sure to extract details mentioned inside boxes and other special formats. Organize the information in the following JSON format:

#     {
#       "Name": "Extract the full name of the applicant.",
#       "Role": "Extract the current or most recent role of the applicant.",
#       "Email": "Extract the applicant's email address.",
#       "Profile": "Extract the applicant's complete brief profile summary or objective.",
#       "Professional Experience": [
#         {
#           "Duration": "Extract the duration of employment in the format 'Month Year - Month Year' or similar.",
#           "Organization": "Extract the name of the organization.",
#           "Designation": "Extract the job title or designation held."
#         }
#       ],
#       "Technical Skills": [
#         {
#           "Skills": "Extract the technical skills listed by the applicant.",
#           "Category": "Categorize the skills if possible, otherwise leave as 'General'."
#         }
#       ],
#       "Certifications": "Extract any certifications mentioned by the applicant.",
#       "Project Experience - Detail": [
#         {
#           "Project Type": "Extract the type of project (e.g., web development, data analysis, etc.).",
#           "Description": "Extract a brief description of the project.",
#           "Role": "Extract the role of the applicant in the project.",
#           "Team Size": "Extract the team size involved in the project.",
#           "Duration": "Extract the duration of the project in the format 'Month Year - Month Year' or similar.",
#           "Tools & Technology": "Extract the tools and technologies used in the project.",
#           "Responsibilities": "Extract complete information of the roles andresponsibilities, Synopsis held by the applicant in the project."
#         }
#       ],
#       "Educational Details": [
#         {
#           "Qualification": "Extract the qualification obtained (e.g., B.Sc., M.Sc., etc.).",
#           "University/Board": "Extract the name of the university or board.",
#           "Year": "Extract the year of completion.",
#           "Percentage": "Extract the percentage or CGPA obtained."
#         }
#       ]
#     }

#     Provide the extracted information in JSON format only.
#     '''
#     try:
#         response = openai.ChatCompletion.create(
#             model="gpt-3.5-turbo",
#             messages=[
#                 {"role": "system", "content": prompt},
#                 {"role": "user", "content": resume_data}
#             ],
#             temperature=0.0,
#             max_tokens=1500
#         )

#         # Debugging: Print response type and content
#         print(f"Response type: {type(response)}")
#         print(f"Response content: {response}")

#         if isinstance(response, str):
#             response_content = response.strip()
#         elif isinstance(response, dict) and 'choices' in response:
#             response_content = response['choices'][0]['message']['content'].strip()
#         else:
#             return {"error": "Unexpected response format from OpenAI API"}

#         # Extract JSON content from the response
#         try:
#             start_index = response_content.index("{")
#             end_index = response_content.rindex("}") + 1
#             json_content = response_content[start_index:end_index]
#             data = json.loads(json_content)
#             return data
#         except (ValueError, json.JSONDecodeError) as e:
#             return {"error": "Invalid response format from OpenAI API"}

#     except Exception as e:
#         return {"error": str(e)}  # Return an error message if API call fails

# @app.route('/')
# def index():
#     return render_template('index.html')

# @app.route("/process", methods=["POST"])
# def ats():
#     doc = request.files.get('resume')
#     company = request.form.get('company')
#     if not doc:
#         return jsonify({"error": "No file uploaded."})

#     try:
#         # Determine the file type and read text accordingly
#         file_ext = doc.filename.split('.')[-1].lower()
#         resume_text = ""

#         if file_ext == 'pdf':
#             # Read text from the PDF in memory
#             reader = PdfReader(doc)
#             for page in reader.pages:
#                 resume_text += page.extract_text()
#         elif file_ext == 'docx':
#             # Read text from the DOCX in memory
#             document = Document(doc)
#             for paragraph in document.paragraphs:
#                 resume_text += paragraph.text + "\n"
#         elif file_ext == 'doc':
#             # Save the uploaded .doc file to a temporary location
#             with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as temp_file:
#                 temp_file.write(doc.read())
#                 temp_file_path = temp_file.name

#             # Initialize COM
#             pythoncom.CoInitialize()
#             try:
#                 # Read text from the DOC file
#                 word = win32com.client.Dispatch("Word.Application")
#                 doc = word.Documents.Open(temp_file_path)
#                 resume_text = doc.Content.Text
#                 doc.Close(False)
#                 word.Quit()
#             finally:
#                 # Uninitialize COM
#                 pythoncom.CoUninitialize()
#                 # Clean up temporary file
#                 os.remove(temp_file_path)
#         else:
#             return jsonify({"error": "Unsupported file type. Please upload a PDF, DOCX, or DOC file."})

#         # Extract ATS data using OpenAI API
#         extracted_data = ats_extractor(resume_text)

#         # Remove fields based on the selected company
#         if company == "bosch":
#             extracted_data.pop("Educational Details", None)
#             extracted_data.pop("Professional Experience", None)

#         # Return extracted data as JSON
#         return jsonify(extracted_data)

#     except Exception as e:
#         return jsonify({"error": str(e)})

# @app.route('/signup', methods=['GET', 'POST'])
# def signup():
#     if request.method == 'POST':
#         full_name = request.form.get('full_name')
#         email = request.form.get('email')
#         password = request.form.get('password')
#         re_password = request.form.get('re_password')

#         # Debugging statements
#         print(f"Full Name: {full_name}")
#         print(f"Email: {email}")
#         print(f"Password: {password}")
#         print(f"Re-entered Password: {re_password}")

#         if password != re_password:
#             flash('Passwords do not match!', 'error')
#             return render_template('index.html', signup_active=True)

#         # Check if the email already exists
#         user = User.query.filter_by(email=email).first()
#         if user:
#             flash('Email address already exists! Please log in.', 'error')
#             return redirect(url_for('signup'))

#         # Create new user
#         new_user = User(full_name=full_name, email=email, password=password)
#         db.session.add(new_user)
#         db.session.commit()

#         flash('Signup successful! Please log in.')
#         return redirect(url_for('login'))

#     return render_template('index.html', signup_active=True)

# @app.route('/login', methods=['GET', 'POST'])
# def login():
#     if request.method == 'POST':
#         email = request.form.get('email')
#         password = request.form.get('password')

#         # Authenticate user
#         user = User.query.filter_by(email=email).first()
#         if not user:
#             flash('Email address not found! Please sign up.', 'error')
#             return render_template('index.html', login_active=True, email=email)

#         if not check_password_hash(user.password, password):
#             flash('Incorrect password!', 'error')
#             return render_template('index.html', login_active=True, email=email, password=password)

#         flash('Login successful!')
#         return redirect(url_for('loading'))

#     return render_template('index.html', login_active=True)


# @app.route('/loading')
# def loading():
#     return render_template('welcome.html')


# # @app.route('/welcome')
# # def welcome():
# #     return render_template('welcome.html')


# @app.route('/dashboard')
# def dashboard():
#     return render_template('home.html')

# if __name__ == '__main__':
#     with app.app_context():
#         db.create_all()
#     app.run(port=5000, debug=True)

from flask import Flask, render_template, request, jsonify
from pypdf import PdfReader
from docx import Document
import openai
import json
import win32com.client  # Only needed if running on Windows for .doc files
import pythoncom  # For COM initialization
import tempfile
import os

app = Flask(__name__)

# Set your OpenAI API key
openai.api_key = 'sk-proj-Cg4Hr5OcghsUvXNpx2GbT3BlbkFJgTsZvax58TeenYEAXk2O'

# Define your ats_extractor function using the OpenAI API
def ats_extractor(resume_data, company):
    prompt = f'''
    You are an expert in resume parsing. Please extract all relevant information from the provided resume file. Make sure to extract details mentioned inside boxes and other special formats. Organize the information in the following JSON format:

    {{
      "Name": "Extract the full name of the applicant.",
      "Role": "Extract the current or most recent role of the applicant.",
      "Email": "Extract the applicant's email address.",
      "Profile": "Extract the applicant's complete brief profile summary and objective.",
      "Professional Experience": [
        {{
          "Duration": "Extract the duration of employment in the format 'Month Year - Month Year' or similar.",
          "Organization": "Extract the name of the organization.",
          "Designation": "Extract the job title or designation held."
        }}
      ],
      "Technical Skills": [
        {{
          "Skills": "Extract the technical skills listed by the applicant.",
          "Category": "Categorize the skills if possible."
        }}
      ],
      "Certifications": "Extract any certifications mentioned by the applicant.",
      "Project Experience - Detail": [
        {{
          "Project Type": "Extract the type of project (e.g., development, analysis, integration, implementation, testing, support etc.).",
          "Description": "Extract a brief description of the project.",
          "Role": "Extract the role of the applicant in the project.",
          "Team Size": "Extract the team size involved in the project.",
          "Duration": "Extract the duration of the project in the format 'Month Year - Month Year' or similar.",
          "Tools & Technology": "Extract the tools and technologies used in the project.",
          "Responsibilities": "Extract complete information of the roles and responsibilities held by the applicant in the project."
        }}
      ],
      "Educational Details": [
        {{
          "Qualification": "Extract the qualification obtained (e.g., B.Sc., M.Sc., etc.).",
          "University/Board": "Extract the name of the university or board.",
          "Year": "Extract the year of completion.",
          "Percentage": "Extract the percentage or CGPA obtained."
        }}
      ]
    }}

    Provide the extracted information in JSON format only.
    '''
    
    if company == "deloitte":
        deloitte_prompt = '''
        Organize the information in the following JSON format for Deloitte:

        {
          "personal_info": {
            "full_name": "Extract the full name of the applicant.",
            "role_designation": "Extract the current or most recent role of the applicant."
          },
          "profile": "Extract the applicant's complete brief profile summary and objective.",
          "skills": {
            "key_skills": ["Extract the key skills listed by the applicant."]
          },
          "education": "Extract the educational details of the applicant in a single line as follows: Name of Examination, Board/University, Name of Institution, Year of Passing, Aggregate.",



          "experience": ["Extract the professional experience of the applicant."],
          "projects": [
            {
              "project_name": "Extract the project name.",
              "client_name": "Extract the client name.",
              "duration": "Extract the project duration in the format 'Month Year - Month Year' or similar.",
              "role": "Extract the role of the applicant in the project.",
              "environment": "Extract the tools and technologies used in the project.",
              "description": "Extract a brief description of the project."
            }
          ]
        }
        '''
        prompt = deloitte_prompt

    try:
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": resume_data}
            ],
            temperature=0.0,
            max_tokens=8192
        )

        # Debugging: Print response type and content
        print(f"Response type: {type(response)}")
        print(f"Response content: {response}")

        if isinstance(response, str):
            response_content = response.strip()
        elif isinstance(response, dict) and 'choices' in response:
            response_content = response['choices'][0]['message']['content'].strip()
        else:
            return {"error": "Unexpected response format from OpenAI API"}

        # Extract JSON content from the response
        try:
            start_index = response_content.index("{")
            end_index = response_content.rindex("}") + 1
            json_content = response_content[start_index:end_index]
            data = json.loads(json_content)
            return data
        except (ValueError, json.JSONDecodeError) as e:
            return {"error": "Invalid response format from OpenAI API"}

    except Exception as e:
        return {"error": str(e)}  # Return an error message if API call fails
    
@app.route('/')
def home():
    return render_template('home.html')

@app.route('/pit')
def index():
    return render_template('index.html')

@app.route("/process", methods=["POST"])
def ats():
    doc = request.files.get('resume')
    company = request.form.get('company')
    if not doc:
        return jsonify({"error": "No file uploaded."})

    try:
        # Determine the file type and read text accordingly
        file_ext = doc.filename.split('.')[-1].lower()
        resume_text = ""

        if file_ext == 'pdf':
            # Read text from the PDF in memory
            reader = PdfReader(doc)
            for page in reader.pages:
                resume_text += page.extract_text()
        elif file_ext == 'docx':
            # Read text from the DOCX in memory
            document = Document(doc)
            for paragraph in document.paragraphs:
                resume_text += paragraph.text + "\n"
        elif file_ext == 'doc':
            # Save the uploaded .doc file to a temporary location
            with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as temp_file:
                temp_file.write(doc.read())
                temp_file_path = temp_file.name

            # Initialize COM
            pythoncom.CoInitialize()
            try:
                # Read text from the DOC file
                word = win32com.client.Dispatch("Word.Application")
                doc = word.Documents.Open(temp_file_path)
                resume_text = doc.Content.Text
                doc.Close(False)
                word.Quit()
            finally:
                # Uninitialize COM
                pythoncom.CoUninitialize()
                # Clean up temporary file
                os.remove(temp_file_path)
        else:
            return jsonify({"error": "Unsupported file type. Please upload a PDF, DOCX, or DOC file."})

        # Extract ATS data using OpenAI API
        extracted_data = ats_extractor(resume_text, company)

        # Remove fields based on the selected company
        if company == "bosch":
            extracted_data.pop("Educational Details", None)
            extracted_data.pop("Professional Experience", None)

        if company == "deloitte":
             extracted_data.pop("Certifications", None)

        # Return extracted data as JSON
        return jsonify(extracted_data)

    except Exception as e:
        return jsonify({"error": str(e)})

if __name__ == "__main__":
   app.run(debug=False , host='0.0.0.0')



# from flask import Flask, render_template, request, jsonify
# from pypdf import PdfReader
# from docx import Document
# import g4f
# import json
# import win32com.client  # Only needed if running on Windows for .doc files
# import pythoncom  # For COM initialization
# import tempfile
# import os

# app = Flask(__name__)

# CONFIG_PATH = r"config.yaml"

# # Define your ats_extractor function using the g4f API
# def ats_extractor(resume_data):
#     prompt = '''
#     You are an expert in resume parsing. Please extract all relevant information from the provided resume file. Make sure to extract details mentioned inside boxes and other special formats. Organize the information in the following JSON format:

#     {
#       "Name": "Extract the full name of the applicant.",
#       "Role": "Extract the current or most recent role of the applicant.",
#       "Email": "Extract the applicant's email address.",
#       "Profile": "Extract the applicant's complete brief profile summary or objective.",
#       "Professional Experience": [
#         {
#           "Duration": "Extract the duration of employment in the format 'Month Year - Month Year' or similar.",
#           "Organization": "Extract the name of the organization.",
#           "Designation": "Extract the job title or designation held."
#         }
#       ],
#       "Technical Skills": [
#         {
#           "Skills": "Extract the technical skills listed by the applicant.",
#           "Category": "Categorize the skills if possible."
#         }
#       ],
#       "Certifications": "Extract any certifications mentioned by the applicant.",
#       "Project Experience - Detail": [
#         {
#           "Project Type": "Extract the type of project (e.g., web development, data analysis, etc.).",
#           "Description": "Extract a brief description of the project.",
#           "Role": "Extract the role of the applicant in the project.",
#           "Team Size": "Extract the team size involved in the project.",
#           "Duration": "Extract the duration of the project in the format 'Month Year - Month Year' or similar.",
#           "Tools & Technology": "Extract the tools and technologies used in the project.",
#           "Responsibilities": "Extract complete information of the roles andresponsibilities, Synopsis held by the applicant in the project."
#         }
#       ],
#       "Educational Details": [
#         {
#           "Qualification": "Extract the qualification obtained (e.g., B.Sc., M.Sc., etc.).",
#           "University/Board": "Extract the name of the university or board.",
#           "Year": "Extract the year of completion.",
#           "Percentage": "Extract the percentage or CGPA obtained."
#         }
#       ]
#     }

#     Provide the extracted information in JSON format only.
#     .
#     '''

#     try:
#         response = g4f.ChatCompletion.create(
#             model="gpt-3.5-turbo",
#             messages=[
#                 {"role": "system", "content": prompt},
#                 {"role": "user", "content": resume_data}
#             ],
#             temperature=0.0,
#             max_tokens=4096
#         )

#         # Debugging: Print response type and content
#         print(f"Response type: {type(response)}")
#         print(f"Response content: {response}")

#         if isinstance(response, str):
#             response_content = response.strip()
#         elif isinstance(response, dict) and 'choices' in response:
#             response_content = response['choices'][0]['message']['content'].strip()
#         else:
#             return {"error": "Unexpected response format from g4f API"}

#         # Extract JSON content from the response
#         try:
#             start_index = response_content.index("{")
#             end_index = response_content.rindex("}") + 1
#             json_content = response_content[start_index:end_index]
#             data = json.loads(json_content)
#             return data
#         except (ValueError, json.JSONDecodeError) as e:
#             return {"error": "Invalid response format from g4f API"}

#     except Exception as e:
#         return {"error": str(e)}  # Return an error message if API call fails

# @app.route('/')
# def index():
#     return render_template('index.html')

# @app.route("/process", methods=["POST"])
# def ats():
#     doc = request.files.get('resume')
#     company = request.form.get('company')
#     if not doc:
#         return jsonify({"error": "No file uploaded."})

#     try:
#         # Determine the file type and read text accordingly
#         file_ext = doc.filename.split('.')[-1].lower()
#         resume_text = ""

#         if file_ext == 'pdf':
#             # Read text from the PDF in memory
#             reader = PdfReader(doc)
#             for page in reader.pages:
#                 resume_text += page.extract_text()
#         elif file_ext == 'docx':
#             # Read text from the DOCX in memory
#             document = Document(doc)
#             for paragraph in document.paragraphs:
#                 resume_text += paragraph.text + "\n"
#         elif file_ext == 'doc':
#             # Save the uploaded .doc file to a temporary location
#             with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as temp_file:
#                 temp_file.write(doc.read())
#                 temp_file_path = temp_file.name

#             # Initialize COM
#             pythoncom.CoInitialize()
#             try:
#                 # Read text from the DOC file
#                 word = win32com.client.Dispatch("Word.Application")
#                 doc = word.Documents.Open(temp_file_path)
#                 resume_text = doc.Content.Text
#                 doc.Close(False)
#                 word.Quit()
#             finally:
#                 # Uninitialize COM
#                 pythoncom.CoUninitialize()
#                 # Clean up temporary file
#                 os.remove(temp_file_path)
#         else:
#             return jsonify({"error": "Unsupported file type. Please upload a PDF, DOCX, or DOC file."})

#         # Extract ATS data using g4f API
#         extracted_data = ats_extractor(resume_text)

#         # Remove fields based on the selected company
#         if company == "bosch":
#             extracted_data.pop("Educational Details", None)
#             extracted_data.pop("Professional Experience", None)

#         # Return extracted data as JSON
#         return jsonify(extracted_data)

#     except Exception as e:
#         return jsonify({"error": str(e)})

# if __name__ == "__main__":
#     app.run(port=8000, debug=True)
